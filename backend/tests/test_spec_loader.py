"""
test_spec_loader.py
spec_loader モジュールのユニットテスト
"""

import os
import pytest
from pathlib import Path
from docx import Document

from spec_loader import load_docx, load_html, load_all_specs, build_system_prompt


@pytest.fixture
def tmp_specs_dir(tmp_path):
    """テスト用の一時仕様書ディレクトリを作成する"""
    return tmp_path


@pytest.fixture
def sample_docx(tmp_specs_dir):
    """テスト用 .docx ファイルを作成する"""
    filepath = tmp_specs_dir / "test.docx"
    doc = Document()
    doc.add_paragraph("テスト仕様書")
    doc.add_paragraph("機能A: ユーザーはログインできる")
    doc.add_paragraph("")  # 空行（除外対象）
    doc.add_paragraph("機能B: ユーザーはレポートを作成できる")
    doc.save(str(filepath))
    return str(filepath)


@pytest.fixture
def sample_html(tmp_specs_dir):
    """テスト用 .html ファイルを作成する"""
    filepath = tmp_specs_dir / "test.html"
    filepath.write_text(
        """<html>
        <head><style>body { color: red; }</style></head>
        <body>
            <h1>画面仕様書</h1>
            <script>alert('test');</script>
            <p>ログイン画面の仕様</p>
        </body>
        </html>""",
        encoding="utf-8",
    )
    return str(filepath)


# ── load_docx ─────────────────────────────────────────────────────────

class TestLoadDocx:
    def test_returns_string(self, sample_docx):
        result = load_docx(sample_docx)
        assert isinstance(result, str)

    def test_extracts_text(self, sample_docx):
        result = load_docx(sample_docx)
        assert "テスト仕様書" in result
        assert "機能A: ユーザーはログインできる" in result
        assert "機能B: ユーザーはレポートを作成できる" in result

    def test_skips_empty_lines(self, sample_docx):
        result = load_docx(sample_docx)
        # 空行は含まれない（strip() で除外されている）
        lines = result.split("\n")
        assert all(line.strip() != "" for line in lines)

    def test_file_not_found_raises(self):
        with pytest.raises(Exception):
            load_docx("/nonexistent/path/file.docx")


# ── load_html ─────────────────────────────────────────────────────────

class TestLoadHtml:
    def test_returns_string(self, sample_html):
        result = load_html(sample_html)
        assert isinstance(result, str)

    def test_extracts_text(self, sample_html):
        result = load_html(sample_html)
        assert "画面仕様書" in result
        assert "ログイン画面の仕様" in result

    def test_removes_script_tags(self, sample_html):
        result = load_html(sample_html)
        assert "alert" not in result

    def test_removes_style_tags(self, sample_html):
        result = load_html(sample_html)
        assert "color: red" not in result

    def test_file_not_found_raises(self):
        with pytest.raises(Exception):
            load_html("/nonexistent/path/file.html")


# ── load_all_specs ────────────────────────────────────────────────────

class TestLoadAllSpecs:
    def test_returns_dict(self, tmp_specs_dir):
        result = load_all_specs(str(tmp_specs_dir))
        assert isinstance(result, dict)

    def test_empty_dir_returns_empty_dict(self, tmp_specs_dir):
        result = load_all_specs(str(tmp_specs_dir))
        assert result == {}

    def test_loads_docx_files(self, tmp_specs_dir, sample_docx):
        result = load_all_specs(str(tmp_specs_dir))
        assert "test.docx" in result
        assert "テスト仕様書" in result["test.docx"]

    def test_loads_html_files(self, tmp_specs_dir, sample_html):
        result = load_all_specs(str(tmp_specs_dir))
        assert "test.html" in result
        assert "画面仕様書" in result["test.html"]

    def test_loads_htm_extension(self, tmp_specs_dir):
        filepath = tmp_specs_dir / "page.htm"
        filepath.write_text("<html><body><p>HTM形式</p></body></html>", encoding="utf-8")
        result = load_all_specs(str(tmp_specs_dir))
        assert "page.htm" in result

    def test_skips_unsupported_extensions(self, tmp_specs_dir):
        (tmp_specs_dir / "readme.txt").write_text("テキストファイル", encoding="utf-8")
        (tmp_specs_dir / "data.pdf").write_bytes(b"%PDF-1.4")
        result = load_all_specs(str(tmp_specs_dir))
        assert "readme.txt" not in result
        assert "data.pdf" not in result

    def test_nonexistent_dir_raises_file_not_found(self):
        with pytest.raises(FileNotFoundError):
            load_all_specs("/nonexistent/specs/dir")

    def test_keys_are_sorted(self, tmp_specs_dir):
        (tmp_specs_dir / "c.html").write_text("<p>C</p>", encoding="utf-8")
        (tmp_specs_dir / "a.html").write_text("<p>A</p>", encoding="utf-8")
        (tmp_specs_dir / "b.html").write_text("<p>B</p>", encoding="utf-8")
        result = load_all_specs(str(tmp_specs_dir))
        assert list(result.keys()) == ["a.html", "b.html", "c.html"]

    def test_loads_multiple_files(self, tmp_specs_dir, sample_docx, sample_html):
        result = load_all_specs(str(tmp_specs_dir))
        assert len(result) == 2
        assert "test.docx" in result
        assert "test.html" in result


# ── build_system_prompt ───────────────────────────────────────────────

class TestBuildSystemPrompt:
    def test_returns_string(self):
        result = build_system_prompt({})
        assert isinstance(result, str)

    def test_empty_specs_returns_prompt(self):
        result = build_system_prompt({})
        assert "FieldReport" in result

    def test_includes_filename(self):
        specs = {"仕様書A.html": "内容A"}
        result = build_system_prompt(specs)
        assert "仕様書A.html" in result

    def test_includes_content(self):
        specs = {"仕様書A.html": "ユーザー認証の仕様"}
        result = build_system_prompt(specs)
        assert "ユーザー認証の仕様" in result

    def test_includes_multiple_specs(self):
        specs = {
            "機能仕様書.docx": "機能Aの説明",
            "画面仕様書.html": "画面Bの説明",
        }
        result = build_system_prompt(specs)
        assert "機能仕様書.docx" in result
        assert "画面仕様書.html" in result
        assert "機能Aの説明" in result
        assert "画面Bの説明" in result

    def test_prompt_contains_instructions(self):
        result = build_system_prompt({})
        assert "仕様書" in result
        assert "矛盾" in result
